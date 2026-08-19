"""Generate .docx analysis report from dashboard stats with embedded charts."""
import io
import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker

# ── Chart style ──
PRIMARY = '#4f46e5'
NON_HR = '#f59e0b'
ACCENT = '#6366f1'
GRAY = '#9ca3af'
LIGHT_BG = '#f9fafb'

plt.rcParams.update({
    'font.family': 'sans-serif',
    'font.size': 9,
    'axes.titlesize': 12,
    'axes.titleweight': 'bold',
    'axes.spines.top': False,
    'axes.spines.right': False,
    'axes.facecolor': 'white',
    'figure.facecolor': 'white',
})


def _chart_to_image(fig):
    """Convert matplotlib figure to BytesIO PNG."""
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf


def _draw_category_bar(stats):
    """Horizontal bar: case type hierarchy (top 10)."""
    types = stats['type_hierarchy'][:10]
    labels = [('Non-HR ' if h['is_non_hr'] else '') + h['major'] for h in types]
    values = [h['count'] for h in types]

    fig, ax = plt.subplots(figsize=(6, 3))
    colors = [NON_HR if h['is_non_hr'] else PRIMARY for h in types]
    bars = ax.barh(labels, values, color=colors, height=0.6)
    ax.invert_yaxis()
    ax.set_xlabel('Case Count')
    ax.set_title('Case Distribution by Category')

    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + max(values) * 0.01, bar.get_y() + bar.get_height() / 2,
                str(val), va='center', fontsize=8, color='#374151')

    ax.xaxis.set_major_locator(ticker.MaxNLocator(integer=True))
    return _chart_to_image(fig)


def _draw_subtype_bar(stats):
    """Horizontal bar: top 15 subtypes across all categories."""
    subtypes = stats.get('subtype_ranking', [])[:15]
    if not subtypes:
        return None

    labels = [s['subtype'] for s in subtypes]
    values = [s['count'] for s in subtypes]

    fig, ax = plt.subplots(figsize=(6, 3.5))
    bars = ax.barh(labels, values, color=PRIMARY, height=0.6)
    ax.invert_yaxis()
    ax.set_xlabel('Case Count')
    ax.set_title('Top Subtypes (Cross-Category)')

    for bar, s in zip(bars, subtypes):
        ax.text(bar.get_width() + max(values) * 0.01, bar.get_y() + bar.get_height() / 2,
                f"{s['count']} ({s['percentage']}%)", va='center', fontsize=7, color='#6b7280')

    ax.xaxis.set_major_locator(ticker.MaxNLocator(integer=True))
    return _chart_to_image(fig)


def _draw_heatmap_chart(stats):
    """Heatmap image: Chat hourly distribution (day x hour)."""
    heatmap_data = stats.get('chat_heatmap', [])
    if not heatmap_data:
        return None

    hours = ['9h', '10h', '11h', '12h', '13h', '14h', '15h', '16h', '17h', '18h']
    dates = [d['date'] for d in heatmap_data]
    matrix = [d['hours'] for d in heatmap_data]

    fig, ax = plt.subplots(figsize=(7, max(3, len(dates) * 0.25)))
    im = ax.imshow(matrix, cmap='Purples', aspect='auto')

    ax.set_xticks(range(len(hours)))
    ax.set_xticklabels(hours, fontsize=7)
    ax.set_yticks(range(len(dates)))
    ax.set_yticklabels(dates, fontsize=7)
    ax.set_title('Chat Incoming Heatmap (Daily x Hourly)')

    # Annotate cells with values > 0
    for i in range(len(dates)):
        for j in range(len(hours)):
            val = matrix[i][j]
            if val > 0:
                ax.text(j, i, str(val), ha='center', va='center',
                        fontsize=6, color='white' if val > max(1, max(matrix[i]) / 2) else '#374151')

    plt.colorbar(im, ax=ax, shrink=0.8, label='Cases')
    return _chart_to_image(fig)


def _draw_hourly_avg_chart(stats):
    """Bar chart: average hourly chat distribution."""
    hourly = stats.get('chat_hourly_avg', [])
    if not hourly:
        return None

    labels = [h['hour'] for h in hourly]
    values = [h['avg'] for h in hourly]

    fig, ax = plt.subplots(figsize=(6, 2.5))
    ax.bar(labels, values, color=PRIMARY, width=0.6, alpha=0.85)
    ax.set_ylabel('Avg Cases')
    ax.set_title('Average Hourly Chat Distribution')

    for i, (label, val) in enumerate(zip(labels, values)):
        if val > 0:
            ax.text(i, val + max(values) * 0.02, str(val),
                    ha='center', fontsize=7, color='#374151')

    ax.yaxis.set_major_locator(ticker.MaxNLocator(integer=True))
    return _chart_to_image(fig)


def _draw_resolution_donut(stats):
    """Donut: resolution breakdown."""
    res = stats.get('resolution_counts', {})
    if not res:
        return None

    labels = list(res.keys())
    values = list(res.values())
    colors = ['#4f46e5', '#9ca3af', '#f87171', '#fb923c'][:len(labels)]

    fig, ax = plt.subplots(figsize=(4, 3))
    wedges, texts, autotexts = ax.pie(
        values, labels=None, autopct='%1.1f%%', startangle=90,
        colors=colors, pctdistance=0.75
    )
    ax.legend(wedges, [f'{l} ({v})' for l, v in zip(labels, values)],
              loc='center left', bbox_to_anchor=(1, 0.5), fontsize=8)
    ax.set_title('Case Resolution')
    return _chart_to_image(fig)


def _draw_daily_trend(stats):
    """Line chart: daily case trend."""
    daily = stats.get('daily_trend', [])
    if not daily:
        return None
    dates = [d['date'] for d in daily]
    counts = [d['count'] for d in daily]

    fig, ax = plt.subplots(figsize=(8, 3))
    ax.fill_between(range(len(dates)), counts, alpha=0.1, color=PRIMARY)
    ax.plot(range(len(dates)), counts, color=PRIMARY, linewidth=1.5, marker='o', markersize=3)
    ax.set_xticks(range(0, len(dates), max(1, len(dates)//15)))
    ax.set_xticklabels([dates[i] for i in range(0, len(dates), max(1, len(dates)//15))], rotation=45, fontsize=7)
    ax.set_title('Daily Case Volume')
    ax.set_ylabel('Cases')
    return _chart_to_image(fig)


# ── Main report generator ─────────────────────────────────────────────────

def generate_report(stats, filename, has_comparison=False):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'PingFang SC')

    # Title
    title = doc.add_heading('HR Case Analysis Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Period: {stats['date_range']['start']} -> {stats['date_range']['end']} "
        f"({stats['date_range']['days']} days)  |  Source: {filename}"
    ).font.size = Pt(9)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)
    doc.add_paragraph()

    # ── 1. KPI ──
    doc.add_heading('1. Key Metrics', level=1)
    kpi_table = doc.add_table(rows=2, cols=7, style='Light Grid Accent 1')
    kpi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sla_str = f"{stats.get('sla_compliance', '—')}%" if stats.get('sla_compliance') else '—'
    kpi_headers = ['Total', 'HR Cases', 'Non-HR', 'Duplicate', '3rd Party', 'SLA Rate', 'CSAT Score']
    kpi_values = [
        str(stats['total_cases']), str(stats['hr_count']), str(stats['non_hr_count']),
        str(stats['duplicate_count']), str(stats['third_party_count']),
        sla_str, f"{stats.get('csat_score', '—')}%" if stats.get('csat_score') else '—',
    ]
    for i, h in enumerate(kpi_headers):
        kpi_table.rows[0].cells[i].text = h
        kpi_table.rows[1].cells[i].text = kpi_values[i]

    total_delta = stats.get('card_comparisons', {}).get('total', {}).get('vs_last')
    if has_comparison and total_delta is not None:
        p = doc.add_paragraph()
        arrow = 'up' if total_delta > 0 else 'down' if total_delta < 0 else 'flat'
        p.add_run(
            f"MoM: {arrow} {abs(total_delta)} cases "
            f"vs prev month"
        ).font.size = Pt(9)

    doc.add_paragraph()

    # ── 2. Category Distribution + Chart ──
    doc.add_heading('2. Case Category Distribution', level=1)
    try:
        chart_img = _draw_category_bar(stats)
        doc.add_picture(chart_img, width=Inches(5.5))
    except Exception:
        doc.add_paragraph('(Chart unavailable)').font.size = Pt(9)

    doc.add_paragraph()
    type_table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
    type_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Category', 'Count', 'Pct', 'MoM', 'Top Subtypes']):
        type_table.rows[0].cells[i].text = h

    for item in stats['type_hierarchy']:
        row = type_table.add_row()
        row.cells[0].text = ('[Non-HR] ' if item['is_non_hr'] else '') + item['major']
        row.cells[1].text = str(item['count'])
        row.cells[2].text = f"{item['percentage']}%"
        if has_comparison and item.get('delta') is not None and item['delta'] != 0:
            a = '+' if item['delta'] > 0 else ''
            row.cells[3].text = f"{a}{item['delta']} ({item.get('delta_pct', '—')}%)"
        else:
            row.cells[3].text = '—'
        minors = [f"{m['name']}({m['count']})" for m in item['minors'][:3]]
        row.cells[4].text = ', '.join(minors)

    doc.add_paragraph()

    # ── 3. Subtype Ranking + Chart ──
    doc.add_heading('3. Top Subtype Issues (Cross-Category)', level=1)
    try:
        subtype_img = _draw_subtype_bar(stats)
        if subtype_img:
            doc.add_picture(subtype_img, width=Inches(5.5))
    except Exception:
        pass

    doc.add_paragraph()
    subtypes = stats.get('subtype_ranking', [])[:15]
    if subtypes:
        st_table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
        st_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(['Subtype', 'Category', 'Count', 'Pct']):
            st_table.rows[0].cells[i].text = h
        for s in subtypes:
            row = st_table.add_row()
            row.cells[0].text = s['subtype']
            row.cells[1].text = s['major']
            row.cells[2].text = str(s['count'])
            row.cells[3].text = f"{s['percentage']}%"

    doc.add_paragraph()

    # ── 4. Chat Heatmap ──
    doc.add_heading('4. Chat Incoming Patterns', level=1)
    try:
        hm_img = _draw_heatmap_chart(stats)
        if hm_img:
            doc.add_picture(hm_img, width=Inches(5.5))
            doc.add_paragraph()
    except Exception:
        pass

    try:
        ha_img = _draw_hourly_avg_chart(stats)
        if ha_img:
            doc.add_picture(ha_img, width=Inches(5))
    except Exception:
        pass

    doc.add_paragraph()

    # ── 5. Daily Trend ──
    daily = stats.get('daily_trend', [])
    if daily:
        doc.add_heading('5. Daily Case Trend', level=1)
        try:
            dt_img = _draw_daily_trend(stats)
            if dt_img:
                doc.add_picture(dt_img, width=Inches(5.5))
        except Exception:
            pass
        doc.add_paragraph()

    # ── 6. Resolution ──
    res = stats.get('resolution_counts', {})
    if res:
        doc.add_heading('6. Resolution Methods', level=1)
        res_table = doc.add_table(rows=1, cols=2, style='Light Grid Accent 1')
        res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(['Resolution', 'Count']):
            res_table.rows[0].cells[i].text = h
        for k, v in sorted(res.items(), key=lambda x: -x[1]):
            row = res_table.add_row()
            row.cells[0].text = k
            row.cells[1].text = str(v)
        doc.add_paragraph()

    # ── 7. Summary ──
    summary = stats.get('monthly_summary', '')
    if summary:
        doc.add_heading('7. Monthly Analysis Summary', level=1)
        for line in summary.split('\n'):
            clean = line.strip().lstrip('#').strip()
            if clean.startswith('---') or not clean:
                continue
            if clean.startswith('*由'):
                p = doc.add_paragraph()
                p.add_run(clean.strip('*')).font.size = Pt(8)
                continue
            if clean.startswith('**') or clean.startswith('###') or clean.startswith('##'):
                doc.add_paragraph(clean.replace('*', '').replace('#', '').strip(), style='List Bullet')
            else:
                p = doc.add_paragraph(clean)
                p.style.font.size = Pt(9)

    # ── Footer ──
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('— Case Intelligence Platform Auto-Generated —').font.size = Pt(8)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
